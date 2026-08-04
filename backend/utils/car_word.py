"""Template WORD (.docx) untuk Corrective Action Report (CAR) — MKS-F-QAD-004.

Alur: Admin upload template .docx berisi placeholder {{key}}. Saat "Cetak PDF",
data CAR disisipkan ke template lalu dikonversi ke PDF via LibreOffice.
Jika belum ada template aktif → fallback ke generator bawaan (reportlab).

Modul ini menyediakan:
- daftar placeholder (CAR_FIELDS) untuk cheatsheet UI,
- pemetaan data CAR → nilai placeholder (car_template_data),
- substitusi placeholder di .docx yang tahan "split runs" (substitute_docx),
- pembuat Starter .docx yang mereplikasi form resmi (build_car_docx_starter),
- render CAR → PDF dari template (render_car_pdf_from_docx).
"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w\.]+)\s*\}\}")

# Placeholder yang tersedia (untuk cheatsheet di UI) + keterangan singkat.
CAR_FIELDS = [
    ("company_name", "Nama perusahaan"),
    ("nc_no", "Nomor CAR"),
    ("status_label", "Status (Open/Assigned/In Progress/Closed)"),
    ("issued_at", "Tanggal terbit"),
    ("issued_by", "Diterbitkan oleh (nama)"),
    ("issued_to", "Ditujukan ke (dept/user)"),
    ("expected_reply_date", "Batas balasan diharapkan"),
    ("chk_inhouse", "Tanda X bila sumber IN-HOUSE"),
    ("chk_external", "Tanda X bila sumber EXTERNAL"),
    ("object_ctx", "Objek NC (Drawing/SO/dll)"),
    ("description", "Deskripsi ketidaksesuaian"),
    ("root_cause", "Root Cause"),
    ("immediate_action", "Immediate Action"),
    ("corrective_action", "Corrective Action"),
    ("preventive_action", "Preventive Action"),
    ("completed_by", "Actions Completed By (nama)"),
    ("completed_date", "Tanggal selesai tindakan"),
    ("dept_head_name", "Approved by Dept. Head (nama)"),
    ("dept_head_date", "Tanggal approve Dept. Head"),
    ("ecn_no", "Nomor ECN terkait"),
    ("initiator_remarks", "Remarks from Initiator"),
    ("chk_risk_yes", "Tanda X bila review risk = Yes"),
    ("chk_risk_no", "Tanda X bila review risk = No"),
    ("effectiveness_reviewed_by", "Effectiveness reviewed by (nama)"),
    ("effectiveness_date", "Tanggal review efektivitas"),
    ("qa_approved_by", "Approved by QA (nama)"),
    ("qa_date", "Tanggal approve QA"),
    ("print_date", "Tanggal cetak"),
    ("printed_by", "Dicetak oleh"),
]

_DEPT_LABEL = {
    "engineering": "Engineering", "qc": "Quality Control", "produksi": "Produksi",
    "sales": "Sales", "purchasing": "Purchasing", "store": "Store",
    "document_control": "Document Control", "finance": "Finance",
    "management": "Management", "other": "Lainnya",
}
_STATUS_LABEL = {"open": "Open", "assigned": "Assigned", "in_progress": "In Progress", "closed": "Closed"}


def _fmt_date(v) -> str:
    if not v:
        return ""
    s = str(v)
    try:
        return datetime.fromisoformat(s.replace("Z", "+00:00")).strftime("%d %b %Y")
    except Exception:
        return s[:10]


def car_template_data(nc: dict, printed_by: str = "") -> dict:
    """Petakan dokumen CAR → nilai placeholder (semua string)."""
    inv = nc.get("investigation") or {}
    clo = nc.get("closeout") or {}
    to_label = _DEPT_LABEL.get(nc.get("issued_to_dept"), nc.get("issued_to") or "")
    if nc.get("issued_to_user"):
        to_label += " · " + (nc["issued_to_user"].get("name") or "")
    if nc.get("link_type") == "drawing" and nc.get("drawing_nos"):
        obj_ctx = "Drawing: " + ", ".join(nc["drawing_nos"])
    else:
        obj_ctx = nc.get("object_ref") or ""
    src = nc.get("source")
    risk = bool(clo.get("risk_review"))
    now = datetime.now(timezone.utc)
    return {
        "company_name": "PT. Mitra Karya Sarana",
        "nc_no": nc.get("nc_no", ""),
        "status": nc.get("status", ""),
        "status_label": _STATUS_LABEL.get(nc.get("status"), nc.get("status") or ""),
        "issued_at": _fmt_date(nc.get("issued_at")),
        "issued_by": (nc.get("issued_by") or {}).get("name", ""),
        "issued_to": to_label,
        "expected_reply_date": _fmt_date(nc.get("expected_reply_date")),
        "chk_inhouse": "X" if src == "in_house" else "",
        "chk_external": "X" if src == "external" else "",
        "object_ctx": obj_ctx,
        "description": nc.get("description") or "",
        "root_cause": inv.get("root_cause") or "",
        "immediate_action": inv.get("immediate_action") or "",
        "corrective_action": inv.get("corrective_action") or "",
        "preventive_action": inv.get("preventive_action") or "",
        "completed_by": inv.get("completed_by") or "",
        "completed_date": _fmt_date(inv.get("completed_date")),
        "dept_head_name": inv.get("dept_head_name") or "",
        "dept_head_date": _fmt_date(inv.get("dept_head_date")),
        "ecn_no": nc.get("ecn_no") or "",
        "initiator_remarks": clo.get("initiator_remarks") or "",
        "chk_risk_yes": "X" if risk else "",
        "chk_risk_no": "" if risk else "X",
        "effectiveness_reviewed_by": clo.get("effectiveness_reviewed_by") or "",
        "effectiveness_date": _fmt_date(clo.get("effectiveness_date")),
        "qa_approved_by": clo.get("qa_approved_by") or "",
        "qa_date": _fmt_date(clo.get("qa_date")),
        "print_date": _fmt_date(now.isoformat()),
        "printed_by": printed_by or "",
    }


def sample_car_data() -> dict:
    """Data contoh untuk preview template."""
    return car_template_data({
        "nc_no": "MKS-QA-CAR-VII-26-01", "status": "in_progress",
        "issued_at": "2026-07-15T00:00:00+00:00", "issued_by": {"name": "Budi (QC)"},
        "issued_to_dept": "engineering", "issued_to_user": {"name": "Adit"},
        "expected_reply_date": "2026-07-22T00:00:00+00:00", "source": "in_house",
        "link_type": "drawing", "drawing_nos": ["DWG.24.12.20_YH.MP1.A.00"],
        "description": "Dimensi lubang baut tidak sesuai gambar (Ø12 seharusnya Ø10).",
        "investigation": {
            "root_cause": "Salah baca revisi gambar.",
            "immediate_action": "Tahan barang, lapor Engineering.",
            "corrective_action": "Revisi drawing & terbitkan ECN.",
            "preventive_action": "Cek revisi terakhir sebelum produksi.",
            "completed_by": "Adit", "completed_date": "2026-07-20T00:00:00+00:00",
            "dept_head_name": "Riski", "dept_head_date": "2026-07-20T00:00:00+00:00",
        },
        "ecn_no": "ECN-26-07-07",
        "closeout": {
            "initiator_remarks": "Tindakan efektif, barang sesuai.",
            "risk_review": True, "effectiveness_reviewed_by": "Budi",
            "effectiveness_date": "2026-07-25T00:00:00+00:00",
            "qa_approved_by": "Salma", "qa_date": "2026-07-26T00:00:00+00:00",
        },
    }, printed_by="admin")


# ── Substitusi placeholder di .docx (tahan split runs) ───────────────────────
def _replace_in_paragraph(paragraph, data: dict) -> None:
    full = "".join(run.text for run in paragraph.runs)
    if "{{" not in full:
        return
    new = PLACEHOLDER_RE.sub(lambda m: str(data.get(m.group(1).strip(), m.group(0))), full)
    if new == full:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nt in cell.tables:
                yield from _iter_table_paragraphs(nt)


def _iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        yield from _iter_table_paragraphs(t)
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                yield p
            for t in hf.tables:
                yield from _iter_table_paragraphs(t)


def substitute_docx(docx_bytes: bytes, data: dict) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))
    for p in _iter_all_paragraphs(doc):
        _replace_in_paragraph(p, data)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def render_car_pdf_from_docx(docx_bytes: bytes, nc: dict, printed_by: str = "") -> bytes:
    """Isi data CAR ke template .docx lalu konversi ke PDF (LibreOffice)."""
    data = car_template_data(nc, printed_by=printed_by)
    sub = substitute_docx(docx_bytes, data)
    from utils.office_render import office_to_pdf
    return office_to_pdf(sub, "docx")


# ── Starter .docx (replika MKS-F-QAD-004 dengan placeholder) ──────────────────
def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcPr.append(shd)


def _p(cell, text, bold=False, size=9, italic=False, align=None, color=None, clear=True):
    para = cell.paragraphs[0] if (clear and cell.paragraphs) else cell.add_paragraph()
    if clear:
        para.text = ""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def build_car_docx_starter() -> bytes:
    """Buat template .docx bawaan yang mereplikasi form MKS-F-QAD-004 + placeholder."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(28)
        section.left_margin = section.right_margin = Pt(28)

    # ── HEADER (2 kolom) ──
    head = doc.add_table(rows=1, cols=2)
    head.style = "Table Grid"
    head.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc, rc = head.rows[0].cells
    lc.width = Pt(200)
    _p(lc, "PT. MITRA KARYA SARANA", bold=True, size=10, align="center")
    _p(lc, "Corrective Action Report (CAR)", bold=True, size=11, align="center", clear=False)
    _p(lc, "CAR No : {{nc_no}}", bold=True, size=9, clear=False)
    # kanan: grid info
    rgrid = rc.add_table(rows=4, cols=1)
    rgrid.style = "Table Grid"
    _p(rgrid.rows[0].cells[0], "Date of Issue : {{issued_at}}", bold=True, size=9)
    _p(rgrid.rows[1].cells[0], "Issued by : {{issued_by}}          Sign & Date :", bold=True, size=9)
    _p(rgrid.rows[2].cells[0], "Issued to : {{issued_to}}", bold=True, size=9)
    _p(rgrid.rows[3].cells[0], "Expected reply date : {{expected_reply_date}}", bold=True, size=9)

    def section_bar(title):
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        c = t.rows[0].cells[0]
        _set_cell_bg(c, "D3D3D3")
        _p(c, title, bold=True, size=9)

    # ── SECTION 1 ──
    section_bar("NONCONFORMANCE INFORMATION (Completed by CAR Initiator)")
    s1 = doc.add_table(rows=1, cols=2)
    s1.style = "Table Grid"
    dcell, chk = s1.rows[0].cells
    dcell.width = Pt(360)
    _p(dcell, "Description of Nonconformance :", bold=True, size=9)
    _p(dcell, "Objek: {{object_ctx}}", italic=True, size=8.5, clear=False)
    _p(dcell, "{{description}}", size=9, clear=False)
    for _ in range(3):
        dcell.add_paragraph("")
    _p(chk, "[{{chk_inhouse}}] IN-HOUSE", size=9)
    _p(chk, "[{{chk_external}}] EXTERNAL", size=9, clear=False)

    # ── SECTION 2 ──
    section_bar("INVESTIGATION & ACTION PLANS (Completed by Responsible Dept./Assignee)")
    s2 = doc.add_table(rows=1, cols=1)
    s2.style = "Table Grid"
    b = s2.rows[0].cells[0]
    _p(b, "Root Cause(s) :", bold=True, size=9)
    _p(b, "{{root_cause}}", size=9, clear=False)
    _p(b, "Immediate Action(s) Taken :", bold=True, size=9, clear=False)
    _p(b, "{{immediate_action}}", size=9, clear=False)
    _p(b, "Corrective Action(s) to eliminate the root cause(s) of NC :", bold=True, size=9, clear=False)
    _p(b, "{{corrective_action}}", size=9, clear=False)
    _p(b, "Preventive Action :", bold=True, size=9, clear=False)
    _p(b, "{{preventive_action}}", size=9, clear=False)

    sign2 = doc.add_table(rows=2, cols=2)
    sign2.style = "Table Grid"
    _p(sign2.rows[0].cells[0], "Actions Completed By / Date :", bold=True, size=9)
    _p(sign2.rows[1].cells[0], "{{completed_by}} / {{completed_date}}", size=9)
    _p(sign2.rows[0].cells[1], "Approved by Dept. Head / Date :", bold=True, size=9)
    _p(sign2.rows[1].cells[1], "{{dept_head_name}} / {{dept_head_date}}", size=9)

    # ── SECTION 3 ──
    section_bar("CAR CLOSEOUT INFORMATION (Completed by Initiator or MR)")
    s3 = doc.add_table(rows=1, cols=1)
    s3.style = "Table Grid"
    c3 = s3.rows[0].cells[0]
    _p(c3, "Remarks from Initiator :", bold=True, size=9)
    _p(c3, "{{initiator_remarks}}", size=9, clear=False)
    _p(c3, "Review of risks and opportunities assessment :  [{{chk_risk_yes}}] Yes   [{{chk_risk_no}}] No   (if yes please attached)",
       bold=True, size=9, clear=False)

    sign3 = doc.add_table(rows=2, cols=2)
    sign3.style = "Table Grid"
    _p(sign3.rows[0].cells[0], "Effectiveness of Actions reviewed by Initiator / Date :", bold=True, size=9)
    _p(sign3.rows[1].cells[0], "{{effectiveness_reviewed_by}} / {{effectiveness_date}}", size=9)
    _p(sign3.rows[0].cells[1], "Approved by QA / Date :", bold=True, size=9)
    _p(sign3.rows[1].cells[1], "{{qa_approved_by}} / {{qa_date}}", size=9)

    foot = doc.add_paragraph()
    fr = foot.add_run("MKS-F-QAD-004#Rev.02          Status: {{status_label}}")
    fr.font.size = Pt(7.5)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
